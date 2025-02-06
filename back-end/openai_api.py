from flask import Flask, request, send_file
from io import BytesIO
import openai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from dotenv import load_dotenv

load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

app = Flask(__name__)

def generate_text(prompt, max_tokens=500):
    """Fetch structured content from OpenAI API."""
    response = openai.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
        max_tokens=max_tokens
    )
    return response.choices[0].message.content

def create_word_document(content) -> BytesIO:
    """Generate a Word document in memory and return buffer."""
    doc = Document()
    
    sections = content.split("\n\n")
    
    for section in sections:
        if section.startswith("# "):
            doc.add_heading(section[2:], level=1)
        elif section.startswith("## "):
            doc.add_heading(section[3:], level=2)
        elif "|" in section:
            rows = section.strip().split("\n")
            num_columns = len(rows[0].split("|")) - 2
            table = doc.add_table(rows=1, cols=num_columns)
            table.style = "Table Grid"
            
            hdr_cells = table.rows[0].cells
            headers = rows[0].split("|")[1:-1]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header.strip()

            for row in rows[2:]:
                cells = row.split("|")[1:-1]
                row_cells = table.add_row().cells
                for i, cell in enumerate(cells):
                    row_cells[i].text = cell.strip()
        else:
            doc.add_paragraph(section)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

